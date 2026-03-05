import streamlit as st
from openai import OpenAI
from PyPDF2 import PdfReader
from docx import Document
import io
import os

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=st.secrets["OPENROUTER_API_KEY"],
)

# --- ИНИЦИАЛИЗАЦИЯ СОСТОЯНИЯ (Кэш и поля) ---
if "input_content" not in st.session_state:
    st.session_state["input_content"] = ""
if "file_key" not in st.session_state:
    st.session_state["file_key"] = 0

# --- ФУНКЦИИ ПАРСИНГА (Должны быть вверху) ---
def extract_text_from_pdf(file):
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        return f"Ошибка при чтении PDF: {e}"

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Ошибка при чтении DOCX: {e}"

def load_law_context():
    laws_text = ""
    # Мы ищем .txt файлы, так как они надежнее для ИИ
    for law_file in ["44fz.txt", "pp2604.txt"]:
        if os.path.exists(law_file):
            try:
                with open(law_file, "r", encoding="utf-8") as f:
                    laws_text += f"\n[ДАННЫЕ ИЗ {law_file}]:\n" + f.read()
            except:
                continue
    return laws_text
    
def reset_app():
    st.session_state["main_text_area"] = "" 
    st.session_state["input_content"] = "" # Было input_text, стало content
    st.session_state["file_key"] += 1
    st.cache_data.clear()
    st.rerun()

def create_docx(report_text):
    doc = Document()
    doc.add_heading('Отчет об аудите критериев', 0)
    doc.add_paragraph(report_text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- КОНФИГУРАЦИЯ ---
st.set_page_config(page_title="Тендерный Аналитик 2604", layout="wide")

# ... (функции extract_text_from_pdf, load_law_context, create_docx остаются без изменений) ...

# --- САЙДБАР ---
with st.sidebar:
    st.title("🔐 Доступ")
    password = st.text_input("Введите пароль", type="password")
    if password != st.secrets["APP_PASSWORD"]:
        st.stop()
    
    st.divider()
    # КНОПКА ОЧИСТКИ (в сайдбаре или внизу)
    if st.button("🗑️ ОЧИСТИТЬ ВСЕ", use_container_width=True, type="primary"):
        reset_app()
    st.caption("Удаляет текст, файлы и сбрасывает память ИИ")

# --- ИНТЕРФЕЙС ---
st.title("🚀 Анализ критериев ЕИС на законность")

col1, col2 = st.columns([1, 1])

with col1:
    option = st.radio("Способ загрузки:", ("Текст", "Документ (PDF/Docx)"))
    
    if option == "Текст":
        # Убираем параметр value, используем только key для прямой связи с session_state
        input_text = st.text_area(
            "Вставьте текст критериев здесь...", 
            height=400,
            key="main_text_area"
        )
        # Синхронизируем для логики очистки
        st.session_state["input_content"] = input_text
    else:
        # file_key меняется при нажатии "Очистить", и загрузчик обнуляется
        uploaded_files = st.file_uploader(
            "Загрузите файлы с критериями (до 10 шт.)", 
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key=f"uploader_{st.session_state['file_key']}"
        )
        
        input_text = ""
        
        if uploaded_files:
            if len(uploaded_files) > 10:
                st.error("Максимальное количество файлов — 10. Пожалуйста, удалите лишние.")
            else:
                combined_text = []
                for file in uploaded_files:
                    if file.name.endswith(".pdf"):
                        text = extract_text_from_pdf(file)
                    else:
                        text = extract_text_from_docx(file)
                    
                    # Добавляем разделитель, чтобы ИИ понимал, где кончается один файл и начинается другой
                    combined_text.append(f"--- СОДЕРЖИМОЕ ФАЙЛА {file.name} ---\n{text}")
                
                input_text = "\n\n".join(combined_text)
                st.success(f"Текст извлечен из {len(uploaded_files)} файл(ов)!")
with col2:
    st.subheader("Результат анализа")
    if st.button("⚖️ Проверить на нарушения", use_container_width=True):
        if not input_text:
            st.warning("Сначала введите данные!")
        else:
            with st.spinner("OpenRouter связывается с Gemini и сверяет законы..."):
                try:
                    legal_context = load_law_context()
                    
                    # --- ТВОЙ ПРОМТ (БЕЗ ИЗМЕНЕНИЙ) ---
                    prompt = f"""
                    Ты — эксперт по извлечению и реструктуризации данных из документов.

                    Тебе будет предоставлено НЕСКОЛЬКО документов одновременно. Каждый документ обозначен маркером:
                    === ДОКУМЕНТ [N] НАЧАЛО ===
                    ...содержимое...
                    === ДОКУМЕНТ [N] КОНЕЦ ===
                    
                    Твоя задача — обработать КАЖДЫЙ документ ОТДЕЛЬНО и независимо друг от друга, не смешивая данные между документами.
                    
                    ---
                    
                    ### СТРУКТУРА ВЫВОДА
                    
                    Для каждого документа выводи блок строго в следующем формате:
                    
                    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                    ## ДОКУМЕНТ [N]
                    ### Заказчик: [Полное наименование заказчика из документа]
                    ### Объект закупки: [Наименование объекта закупки из документа]
                    ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
                    
                    Затем следует полная структурированная обработка документа по правилам ниже.
                    После окончания документа — вставь разделитель: ════════════════════════════════════════
                    
                    ---
                    
                    ### ПРАВИЛА ПАРСИНГА (применять к каждому документу)
                    
                    1. **Таблицы → текст**
                       - Каждую таблицу преобразуй в именованный раздел с подзаголовком.
                       - Каждую строку таблицы превращай в отдельный абзац или пункт вида: «[Название столбца]: [Значение]».
                       - Если строка содержит несколько значимых полей — перечисляй через точку с запятой или выноси в подпункты.
                       - Если таблица описывает критерии/шкалы/формулы — сохраняй логику и числа точно, представляй в виде нумерованных или маркированных пунктов.
                    
                    2. **Реквизиты и коды**
                       - Все ИНН, КПП, ОКТМО, адреса, телефоны, email — выноси в раздел «Реквизиты заказчика» в виде пар «ключ: значение».
                    
                    3. **Структура документа**
                       - Сохраняй исходную иерархию: разделы (I, II, III...) → подразделы (1.1, 1.2...) → пункты → подпункты.
                       - Используй заголовки: ## для разделов, ### для подразделов, #### для пунктов.
                       - Нумерацию и названия разделов бери из оригинала без изменений.
                    
                    4. **Формулы и шкалы**
                       - Формулы сохраняй в исходном виде.
                       - Шкалы баллов оформляй как список: «0 баллов — [условие]», «10 баллов — [условие]» и т.д.
                    
                    5. **Что НЕ делать**
                       - Не пропускай ни одной строки, ни одной ячейки — даже если она кажется технической.
                       - Не переформулируй суть — только структурируй подачу.
                       - Если ячейка пуста — пиши: «не указано».
                       - НИКОГДА не переносить данные из одного документа в другой.
                    
                    6. **Результат**
                       - Единый текст, разбитый на блоки по документам, без таблиц, без ASCII-рамок.
                       - Каждый смысловой блок отделяй пустой строкой.
                       - В конце всех документов выведи сводную таблицу:
                    
                    ---
                    
                    ### СВОДНАЯ ТАБЛИЦА (в конце, после всех документов)
                    
                    | № | Заказчик | Объект закупки | ИНН | Кол-во критериев оценки |
                    |---|----------|----------------|-----|--------------------------|
                    | 1 | ...      | ...            | ... | ...                      |
                    
                    ---
                    
                    ### ДОКУМЕНТЫ ДЛЯ ОБРАБОТКИ:
                    
                    === ДОКУМЕНТ 1 НАЧАЛО ===
                    [ВСТАВЬ ТЕКСТ ДОКУМЕНТА 1]
                    === ДОКУМЕНТ 1 КОНЕЦ ===
                    
                    === ДОКУМЕНТ 2 НАЧАЛО ===
                    [ВСТАВЬ ТЕКСТ ДОКУМЕНТА 2]
                    === ДОКУМЕНТ 2 КОНЕЦ ===
                    
                    ... и так далее до ДОКУМЕНТА N
                    
                    ЭТАЛОННЫЕ ЗАКОНЫ:
                    {legal_context}
                    
                    КРИТЕРИИ ЗАКАЗЧИКА:
                    {input_text}
                    
                    ИНСТРУКЦИЯ:
                    1. В самом начале напиши жирным шрифтом: "СТАТУС: ЗАМЕЧАНИЯ ВЫЯВЛЕНЫ" или "СТАТУС: ЗАМЕЧАНИЙ НЕ ВЫЯВЛЕНО".
                    2. Если нарушения есть, перечисли их по пунктам:
                       - Суть нарушения.
                       - Ссылка на статью (44-ФЗ или ПП 2604).
                       - Текст для жалобы в ФАС.
                    3. Если нарушений нет, кратко подтверди соответствие.
                    
                    Отвечай на русском языке.
                    """
                    
                    # Исправленный вызов API
                    response = client.chat.completions.create(
                        model="anthropic/claude-3.5-sonnet",
                        messages=[{"role": "user", "content": prompt}]
                    )
                    
                    report_content = response.choices[0].message.content
                    
                    # 1. Вывод отчета в окно
                    st.markdown(report_content)
                    
                    # 2. Создание и кнопка скачивания DOCX
                    docx_file = create_docx(report_content)
                    st.divider()
                    st.download_button(
                        label="📄 Скачать отчет в Word (.docx)",
                        data=docx_file,
                        file_name="Audit_Report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Произошла ошибка: {e}")

st.divider()
st.caption("Тендерный отдел | Анализ на основе ИИ")
